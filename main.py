# Standard Imports
import logging
import os

# External Imports
import docx

# Local Imports
from data.update_list import data




"""
Base Class for Text phrases.
TODO, setup the ability for text formatting memorization.
As of now, any updated text reformats text to document default
"""
class Text():
    def __init__(self):
        self.alignment = ''
        self.font = ''
        self.font_size = ''
        self.content = ''


    #If phrase exists, Cycle thru the update dictionary to replace the key,value pairs
    def text_check(phrase, dic):
        if phrase == None or phrase == '':
            return phrase

        else:
            for key, value in dic.items():

                # Searching the phrase for the given key, returns the index where the key starts and ends. If no match, start_index is returned as -1
                start_index = phrase.find(key)
                end_index = start_index + len(key)

                if start_index != -1:
                    print(f'MATCH! Changing {key} to {value}')

                    # If the match is at the start of the textline
                    if start_index == 0:
                        splits = phrase.split(key)
                        phrase = value + splits[1]

                    # If the match is at the end of the textline
                    elif end_index == len(phrase):
                        splits = phrase.split(key)
                        phrase = splits[0] + value

                    # If the match is somewhere in the middle
                    else:
                        splits = phrase.split(key)
                        phrase = splits[0] + value + splits[1]
                    
            return phrase




"""
Base Class to work with Word Documents
Methods to Load and save documents
Methods to loop thru entire word documents to update Text Class
TODO, the updating loops are similar for Footer, Header, and Paragraphs. Potential for optimization there.
"""
class Word():
    def __init__(self, file):
        self.old_name = file
        self.new_name = Text.text_check(phrase = self.old_name, dic = data) # Updating file name, if needed
        self.data = data
        self.doc = self.load_doc()


    def load_doc(self):
        self.log_change(content=f'starting transition from {self.old_name} to {self.new_name}')
        return docx.Document(f'input/{self.old_name}')


    def save_doc(self):
        return self.doc.save(f'output/{self.new_name}')


    # Method to update log.txt if any changes were made (useful if errors after script was run)
    def log_change(self, content):
        with open('data/log.txt', 'a') as f:
            f.write(f'\n {content}')


    def header_loop(self):
        i = 0
        while i < len(self.doc.sections[0].header.tables): 
            print(f'\n Checking Header Table {i} \n')

            j = 0
            while j < len(self.doc.sections[0].header.tables[i]._cells):
                text_line = self.doc.sections[0].header.tables[i]._cells[j].text

                if text_line != '' and text_line != None:
                    new_textline = Text.text_check(phrase=text_line, dic=self.data)

                    # Check to see if the new phrase is the exact same as the old phrase
                    # Due to text formatting over-writting issues, It is best to avoid updating text on the document that doesn't need to be updated. 
                    if new_textline != text_line:
                        try:
                            self.doc.sections[0].header.tables[i]._cells[j].text = new_textline
                            self.log_change(content=f'{new_textline} /// on header table {i}, cell {j}') # Log this change
                        except:
                            logging.error('Header could not be looped thru!') # Error handling if docx can't be updated
                            self.log_change(content=f' E R R O R! on header table {i}, cell {j}') # Log this error
                j += 1
            i += 1


    # Loop through the document's Paragraphs
    def paragraph_loop(self):
        i = 0
        while i < len(self.doc.paragraphs):
            print(f'\n Checking Paragraph {i} \n')
            text_line = self.doc.paragraphs[int(i)].text

            if text_line != '' and text_line != None:
                new_textline = Text.text_check(phrase=text_line, dic=self.data)

                # Check to see if the new phrase is the exact same as the old phrase
                # Due to text formatting over-writting issues, It is best to avoid updating text on the document that doesn't need to be updated. 
                if new_textline != text_line:
                    try:
                        self.doc.paragraphs[i].text = new_textline
                        self.log_change(content=f'>>>{new_textline} on paragraph {i}') # Log this change
                    except:
                        logging.error('Paragraph could not be looped thru!') # Error handling if docx can't be updated
                        self.log_change(content=f' E R R O R! on paragraph {i}!') # Log this error
            i += 1


    # Loop through the document's Tables
    def table_loop(self):
        i = 0
        while i < len(self.doc.tables):
            print(f'\n Checking Table {i} \n')

            j = 0
            while j < len(self.doc.tables[i]._cells):
                text_line = self.doc.tables[i]._cells[j].text

                if text_line != '' and text_line != None:
                    new_textline = Text.text_check(phrase=text_line, dic=self.data)

                    # Check to see if the new phrase is the exact same as the old phrase
                    # Due to text formatting over-writting issues, It is best to avoid updating text on the document that doesn't need to be updated. 
                    if new_textline != text_line:
                        try:
                            self.doc.tables[i]._cells[j].text = new_textline
                            self.log_change(content=f'{new_textline} on body table {i}, cell {j}') # Log this error
                        except:
                            logging.error('Table could not be looped thru!') # Error handling if docx can't be updated
                            self.log_change(content=f' E R R O R! on body table {i}, cell {j}') # Log this error
                j += 1
            i += 1




    # Loop through the document's Footers
    def footer_loop(self):
        i = 0
        while i < len(self.doc.sections[0].footer.tables):
            print(f'\n Checking Footer Table {i} \n')

            j = 0
            while j < len(self.doc.sections[0].footer.tables[i]._cells):
                text_line = self.doc.sections[0].footer.tables[i]._cells[j].text

                if text_line != '' and text_line != None:
                    new_textline = Text.text_check(phrase=text_line, dic=self.data)

                    # Check to see if the new phrase is the exact same as the old phrase
                    # Due to text formatting over-writting issues, It is best to avoid updating text on the document that doesn't need to be updated. 
                    if new_textline != text_line:
                        try:
                            self.doc.sections[0].footer.tables[i]._cells[j].text = new_textline # Log this error
                            self.log_change(content=f'{new_textline} on footer table {i}, cell {j}')
                        except:
                            logging.error('Table could not be looped thru!') # Error handling if docx can't be updated
                            self.log_change(content=f' E R R O R! on body table {i}, cell {j}') # Log this error
                j += 1
            i += 1


    # For each document, loop through all sections, with the below methods
    def loop_thru_document(self):
        self.header_loop()
        self.paragraph_loop()
        self.table_loop()
        self.footer_loop()




"""
General Class for looping thru a folder of Word documents
"""
class App():
    def __init__(self):
        self.input_path = './input'
        self.filename_list = self.get_files()

    def get_files(self):
        filelist = []
        dirpath = os.listdir(self.input_path)

        for file in dirpath:
            if os.path.isfile(os.path.join(self.input_path, file)):
                filelist.append(file)

        print(filelist)
        return filelist


    def delete_old_log(self):
        if os.path.exists('data/log.txt'):
            os.remove('data/log.txt')
        else:
            print('Log file not found')


    # Main scripting loop through all the files
    def main_loop(self):
        self.delete_old_log()
        for file in self.filename_list:
            print(f'\n Loading {file} \n')
            word = Word(file)
            word.loop_thru_document()
            word.save_doc()




if __name__ == '__main__':
    app = App()
    app.main_loop()